import re
import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt

# Mappning av Ami Pro escape-koder till Unicode
ESCAPE_MAP = {
    r'<\v>': 'ö', r'<\V>': 'Ö',
    r'<\d>': 'ä', r'<\D>': 'Ä',
    r'<\e>': 'å', r'<\E>': 'Å',
    r'<\i>': 'é',
    r'<\|>': 'ü',
    r'<\`>': 'à',
}

# Mappning av Ami Pro styckeformat till Word-stilar
STYLE_MAP = {
    'Rubrik':   'Heading 1',
    'Rubrik 1': 'Heading 1',
    'Rubrik 2': 'Heading 2',
    'Rubrik 3': 'Heading 3',
    'Bomb':     'List Bullet',
    'Bomb 1':   'List Bullet',
    'Body Text':'Normal',
}

def replace_escapes(text):
    for code, char in ESCAPE_MAP.items():
        text = text.replace(code, char)
    return text

def clean_line(line):
    # Ersätt escape-koder
    line = replace_escapes(line)  # Först!
    # Ta bort formateringskoder som <:f>, <:s>, <:#548,6151> etc
    line = re.sub(r'<:[^>]*>', '', line)
    # Ta bort <*-> och liknande
    line = re.sub(r'<[^>]*>', '', line)
    return line.strip()

def parse_sam(filepath):
    with open(filepath, 'r', encoding='cp1252', errors='replace') as f:
        content = f.read()

    # Hitta innehållet efter [edoc]
    edoc_pos = content.find('[edoc]')
    if edoc_pos == -1:
        print("Varning: [edoc]-tagg saknas, läser hela filen")
        body = content
    else:
        body = content[edoc_pos + len('[edoc]'):]
    # Debug: visa råtext runt första escape-koden
    pos = body.find('<\\')

    paragraphs = []
    for line in body.splitlines():
        line = line.strip()
        if not line:
            continue

        # Hoppa över sidhuvuden
        if line.startswith('@sidhuvud@'):
            continue

        # Hoppa över rader som bara är >
        if line == '>':
            continue

        # Kontrollera om raden börjar med ett styckeformat: @Rubrik@, @Bomb@ etc
        style_match = re.match(r'^@([^@]+)@(.*)', line)
        if style_match:
            style_name = style_match.group(1).strip()
            text = clean_line(style_match.group(2))
            word_style = STYLE_MAP.get(style_name, 'Normal')
            if text:
                paragraphs.append((word_style, text, style_name))
        else:
            text = clean_line(line)
            if text:
                paragraphs.append(('Normal', text, ''))

    return paragraphs

def convert(sam_path):
    sam_path = Path(sam_path)
    out_path = sam_path.with_suffix('.docx')
    out_path = Path(str(out_path).lower())

    print(f"Läser: {sam_path}")
    paragraphs = parse_sam(sam_path)

    doc = Document()

    for style, text, style_name in paragraphs:
        try:
            p = doc.add_paragraph(text, style=style)
        except Exception:
            p = doc.add_paragraph(text)
        
        if style_name in ('Bibeltext', 'Bibeltext2'):
            # Kursiv text
            for run in p.runs:
                run.italic = True
            # Indrag
            from docx.shared import Cm
            p.paragraph_format.left_indent = Cm(1.5)

    doc.save(out_path)
    print(f"Sparad: {out_path}")
    print(f"Antal stycken: {len(paragraphs)}")

import os

def convert_all(folder="."):
    folder = Path(folder)
    seen = set()
    sam_files = []
    for f in folder.iterdir():
        if f.suffix.lower() == '.sam' and f.name.lower() not in seen:
            seen.add(f.name.lower())
            sam_files.append(f)
    
    if not sam_files:
        print("Inga .SAM-filer hittades.")
        return
    
    print(f"Hittade {len(sam_files)} filer.\n")
    
    for sam_path in sam_files:
        try:
            # Spara tidsstämplar
            stat = sam_path.stat()
            mtime = stat.st_mtime
            atime = stat.st_atime

            convert(sam_path)

            # Återställ tidsstämplar på den nya .docx-filen
            out_path = Path(str(sam_path.with_suffix('.docx')).lower())
            os.utime(out_path, (atime, mtime))

        except Exception as e:
            print(f"  ✗ Fel vid konvertering av {sam_path.name}: {e}")

if __name__ == "__main__":
    if len(sys.argv) == 2:
        arg = sys.argv[1]
        if arg.lower().endswith('.sam'):
            convert(sys.argv[1])
        else:
            convert_all(sys.argv[1])
    else:
        convert_all()
